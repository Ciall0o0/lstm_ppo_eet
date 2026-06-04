"""
批量将 Word 文档中的 LaTeX 公式转换为 Word 原生公式格式（OMML）。

用法: python latex_to_word.py input.docx output.docx

支持:
  - $...$ 行内公式
  - $$...$$ 独立行公式（居中显示，带编号）
"""

import re
import sys
import copy
from pathlib import Path

import latex2mathml.converter
from docx import Document
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from lxml import etree


# ── MathML → OMML 转换 ──────────────────────────────────────────────

# OMML XSLT 转换（将 MathML 转为 Word 公式格式）
# 来自 Microsoft Office 的 OMML2MML.XSLT 反向逻辑
MATHML_TO_OMML_XSLT = r"""<xsl:stylesheet version="1.0"
  xmlns:xsl="http://www.w3.org/1999/XSL/Transform"
  xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"
  xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <xsl:output method="xml" indent="no" omit-xml-declaration="yes"/>

  <xsl:template match="/">
    <m:oMath>
      <xsl:apply-templates select="*"/>
    </m:oMath>
  </xsl:template>

  <!-- 根元素 m:math -->
  <xsl:template match="m:math">
    <m:oMath>
      <xsl:apply-templates/>
    </m:oMath>
  </xsl:template>

  <!-- mrow: 包装元素 -->
  <xsl:template match="m:mrow">
    <xsl:apply-templates/>
  </xsl:template>

  <!-- mi: 标识符 -->
  <xsl:template match="m:mi">
    <m:r>
      <m:rPr>
        <m:sty m:val="p"/>
      </m:rPr>
      <m:t><xsl:value-of select="."/></m:t>
    </m:r>
  </xsl:template>

  <!-- mn: 数字 -->
  <xsl:template match="m:mn">
    <m:r>
      <m:t><xsl:value-of select="."/></m:t>
    </m:r>
  </xsl:template>

  <!-- mo: 操作符 -->
  <xsl:template match="m:mo">
    <m:r>
      <m:t><xsl:value-of select="."/></m:t>
    </m:r>
  </xsl:template>

  <!-- mtext: 文本 -->
  <xsl:template match="m:mtext">
    <m:r>
      <m:rPr>
        <m:nor/>
      </m:rPr>
      <m:t><xsl:value-of select="."/></m:t>
    </m:r>
  </xsl:template>

  <!-- msup: 上标 -->
  <xsl:template match="m:msup">
    <m:sSup>
      <m:e><xsl:apply-templates select="*[1]"/></m:e>
      <m:sup><xsl:apply-templates select="*[2]"/></m:sup>
    </m:sSup>
  </xsl:template>

  <!-- msub: 下标 -->
  <xsl:template match="m:msub">
    <m:sSub>
      <m:e><xsl:apply-templates select="*[1]"/></m:e>
      <m:sub><xsl:apply-templates select="*[2]"/></m:sub>
    </m:sSub>
  </xsl:template>

  <!-- msupsub: 上下标 -->
  <xsl:template match="m:msubsup">
    <m:sSubSup>
      <m:e><xsl:apply-templates select="*[1]"/></m:e>
      <m:sub><xsl:apply-templates select="*[2]"/></m:sub>
      <m:sup><xsl:apply-templates select="*[3]"/></m:sup>
    </m:sSubSup>
  </xsl:template>

  <!-- mfrac: 分数 -->
  <xsl:template match="m:mfrac">
    <m:f>
      <m:num><xsl:apply-templates select="*[1]"/></m:num>
      <m:den><xsl:apply-templates select="*[2]"/></m:den>
    </m:f>
  </xsl:template>

  <!-- msqrt: 平方根 -->
  <xsl:template match="m:msqrt">
    <m:rad>
      <m:radPr><m:degHide m:val="1"/></m:radPr>
      <m:deg/>
      <m:e><xsl:apply-templates/></m:e>
    </m:rad>
  </xsl:template>

  <!-- mroot: n次根 -->
  <xsl:template match="m:mroot">
    <m:rad>
      <m:deg><xsl:apply-templates select="*[2]"/></m:deg>
      <m:e><xsl:apply-templates select="*[1]"/></m:e>
    </m:rad>
  </xsl:template>

  <!-- mfenced: 括号 -->
  <xsl:template match="m:mfenced">
    <m:d>
      <m:dPr>
        <m:begChr m:val="{@open}"/>
        <m:endChr m:val="{@close}"/>
      </m:dPr>
      <m:e><xsl:apply-templates/></m:e>
    </m:d>
  </xsl:template>

  <!-- mo with parentheses/brackets as delimiters -->
  <xsl:template match="m:mo[text()='(' or text()==')' or text()='[' or text()==']' or text()='{' or text()='}']">
    <m:d>
      <m:dPr>
        <xsl:choose>
          <xsl:when test="text()='('"><m:begChr m:val="("/><m:endChr m:val=")"/></xsl:when>
          <xsl:when test="text()='['"><m:begChr m:val="["/><m:endChr m:val="]"/></xsl:when>
          <xsl:when test="text()='{'"><m:begChr m:val="{"/><m:endChr m:val="}"/></xsl:when>
          <xsl:otherwise><m:begChr m:val="("/><m:endChr m:val=")"/></xsl:otherwise>
        </xsl:choose>
      </m:dPr>
    </m:d>
  </xsl:template>

  <!-- munder: 下标（如求和下限） -->
  <xsl:template match="m:munder">
    <m:limLow>
      <m:e><xsl:apply-templates select="*[1]"/></m:e>
      <m:lim><xsl:apply-templates select="*[2]"/></m:lim>
    </m:limLow>
  </xsl:template>

  <!-- mover: 上标（如帽子） -->
  <xsl:template match="m:mover">
    <m:limUpp>
      <m:e><xsl:apply-templates select="*[1]"/></m:e>
      <m:lim><xsl:apply-templates select="*[2]"/></m:lim>
    </m:limUpp>
  </xsl:template>

  <!-- munderover: 上下限 -->
  <xsl:template match="m:munderover">
    <m:nary>
      <m:naryPr>
        <m:chr m:val="{*[1]}"/>
      </m:naryPr>
      <m:sub><xsl:apply-templates select="*[2]"/></m:sub>
      <m:sup><xsl:apply-templates select="*[3]"/></m:sup>
      <m:e/>
    </m:nary>
  </xsl:template>

  <!-- 默认：递归处理子元素 -->
  <xsl:template match="*">
    <xsl:apply-templates/>
  </xsl:template>

</xsl:stylesheet>"""


def mathml_to_omml(mathml_str: str) -> OxmlElement:
    """将 MathML 字符串转换为 OMML 元素。"""
    # 解析 MathML
    mathml_tree = etree.fromstring(mathml_str.encode("utf-8"))

    # 应用 XSLT 转换
    xslt_tree = etree.fromstring(MATHML_TO_OMML_XSLT.encode("utf-8"))
    transform = etree.XSLT(xslt_tree)
    omml_tree = transform(mathml_tree)

    # 返回 OMML 元素
    return omml_tree.getroot()


def latex_to_omml(latex_str: str) -> OxmlElement:
    """将 LaTeX 字符串转换为 OMML 元素。"""
    # 转换 LaTeX → MathML
    mathml_str = latex2mathml.converter.convert(latex_str)
    # 转换 MathML → OMML
    return mathml_to_omml(mathml_str)


# ── 文本分割与公式提取 ──────────────────────────────────────────────

def split_text_with_formulas(text: str):
    """
    将文本按 LaTeX 公式分割。
    返回 [(text, is_formula), ...] 列表。
    is_formula=True 表示该段是 LaTeX 公式。
    """
    # 匹配 $$...$$ 和 $...$
    # 优先匹配 $$（贪婪匹配到最近的 $$）
    pattern = r'\$\$(.+?)\$\$|\$([^\$]+?)\$'
    result = []
    last_end = 0

    for match in re.finditer(pattern, text):
        # 公式前的普通文本
        if match.start() > last_end:
            result.append((text[last_end:match.start()], False))

        # 公式内容
        formula = match.group(1) or match.group(2)
        is_display = match.group(1) is not None  # $$...$$ 是独立行
        result.append((formula, True, is_display))
        last_end = match.end()

    # 最后的普通文本
    if last_end < len(text):
        result.append((text[last_end:], False))

    return result


def extract_formula_number(text: str):
    """
    从公式后提取编号，如 "公式 (3-1)" 或 "(3-1)"。
    返回 (formula_text, number_text)。
    """
    match = re.match(r'^(.*?)\s*(\([\d\-]+\))\s*$', text)
    if match:
        return match.group(1).strip(), match.group(2)
    return text, None


# ── 文档处理 ──────────────────────────────────────────────────────

def process_paragraph(paragraph):
    """处理一个段落，将其中的 LaTeX 公式转换为 Word 公式。"""
    full_text = paragraph.text
    if not full_text or '$' not in full_text:
        return False

    # 分割文本
    parts = split_text_with_formulas(full_text)
    if not any(p[1] for p in parts):
        return False

    # 保存原始格式信息
    original_runs = list(paragraph.runs)
    if not original_runs:
        return False

    # 获取第一个 run 的格式作为参考
    ref_run = original_runs[0]
    ref_font_name = ref_run.font.name or "宋体"
    ref_font_size = ref_run.font.size

    # 清除段落中所有现有的 run
    for run in original_runs:
        run._element.getparent().remove(run._element)

    # 重新构建段落内容
    for part in parts:
        if len(part) == 2:
            text, is_formula = part
            is_display = False
        else:
            text, is_formula, is_display = part

        if not text:
            continue

        if is_formula:
            # 提取公式编号
            formula_text, formula_num = extract_formula_number(text)
            try:
                # 转换 LaTeX → OMML
                omml_elem = latex_to_omml(formula_text)
                paragraph._element.append(omml_elem)

                # 如果有编号，添加编号文本
                if formula_num:
                    num_run = OxmlElement("w:r")
                    num_t = OxmlElement("w:t")
                    num_t.text = f"  {formula_num}"
                    num_run.append(num_t)
                    paragraph._element.append(num_run)

            except Exception as e:
                # 转换失败，保留原始 LaTeX 文本
                print(f"  [警告] 公式转换失败: {formula_text[:50]}... ({e})")
                run = paragraph.add_run(f"${text}$")
                run.font.name = ref_font_name
                if ref_font_size:
                    run.font.size = ref_font_size
        else:
            # 普通文本
            run = paragraph.add_run(text)
            run.font.name = ref_font_name
            if ref_font_size:
                run.font.size = ref_font_size

    return True


def process_table_cell(cell):
    """处理表格单元格中的公式。"""
    for paragraph in cell.paragraphs:
        process_paragraph(paragraph)


def convert_latex_in_docx(input_path: str, output_path: str):
    """主函数：批量转换文档中的 LaTeX 公式。"""
    print(f"打开文档: {input_path}")
    doc = Document(input_path)

    converted_count = 0

    # 处理正文段落
    print("处理正文段落...")
    for i, para in enumerate(doc.paragraphs):
        if process_paragraph(para):
            converted_count += 1
            print(f"  段落 {i+1}: 已转换公式")

    # 处理表格
    print("处理表格...")
    for table_idx, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if process_paragraph(para):
                        converted_count += 1

    # 保存
    print(f"保存文档: {output_path}")
    doc.save(output_path)
    print(f"完成！共转换 {converted_count} 处公式。")


# ── 主入口 ──────────────────────────────────────────────────────

if __name__ == "__main__":
    input_file = sys.argv[1] if len(sys.argv) > 1 else "backup.docx"
    output_file = sys.argv[2] if len(sys.argv) > 2 else "backup_converted.docx"

    if not Path(input_file).exists():
        print(f"错误: 文件不存在 - {input_file}")
        sys.exit(1)

    convert_latex_in_docx(input_file, output_file)
